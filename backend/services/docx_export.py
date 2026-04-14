from docx import Document
import uuid

def create_docx(title, content):
    doc = Document()
    doc.add_heading(title, 0)

    for para in content.split("\n"):
        doc.add_paragraph(para)

    file_path = f"essay_{uuid.uuid4().hex}.docx"
    doc.save(file_path)

    return file_path